import sys

from docx import Document
from docx.shared import Pt
from PyQt5 import QtCore, QtWidgets

import source.py_ui.screen as screen
import source.tools.Constants as Constants
import source.tools.Containers as Containers


class mywindow(QtWidgets.QMainWindow, screen.Ui_screen):
    def __init__(self):
        super().__init__()
        self.setupUi(self)

        self.action.triggered.connect(self.generate)
        self.setWindowTitle("speech-therapist-assistant")

        self.tabs = []

        for key in Constants.DATA.keys():
            tab = Containers.Tab(None, Constants.DATA[key], key)
            self.tabWidget.addTab(tab.out, key)
            self.tabs.append(tab)

    def generate(self):
        text = ""
        document = Document()
        run = document.add_paragraph().add_run()
        font = run.font
        font.name = "Times New Roman"
        font.size = Pt(7)
        for tab in self.tabs:
            add = tab.getData()
            text += add + "\n"
            document.add_paragraph(add)

        options = QtWidgets.QFileDialog.Options()
        options |= QtWidgets.QFileDialog.DontUseNativeDialog
        fileName, _ = QtWidgets.QFileDialog.getSaveFileName(
            self,
            caption="Выберите файл для сохранения",
            filter="Документы (.docx)",
            directory="Название",
            options=options,
        )
        if fileName:
            # print(fileName)
            if fileName[::-1][0:5] != ".docx"[::-1]:
                fileName += ".docx"
            document.save(fileName)

    def resizeEvent(self, event):
        # old_size = [event.oldSize().width(), event.oldSize().height()]
        current_size = [event.size().width(), event.size().height()]

        self.tabWidget.setGeometry(0, 0, current_size[0], current_size[1] - 50)


if __name__ == "__main__":
    app = QtWidgets.QApplication([])
    application = mywindow()
    application.show()

    sys.exit(app.exec())
